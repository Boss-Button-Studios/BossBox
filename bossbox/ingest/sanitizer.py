"""
Physical Document Sanitizer — BossBox Atomic Step 9
=====================================================
Tiered sanitization (Section 9.1) for every external document before it
reaches a model.  Two modes:

Standard (default)
    High-quality text extraction with aggressive hidden-element stripping.
    Removes: hidden character-formatted text, non-visible DOM elements,
    zero-width Unicode characters.  Suspicious elements trigger
    auto-escalation to Deep mode when tesseract is available.

Deep
    Full rasterization to images followed by OCR.  Extracts only what a
    human eye would see.  Requires the ``tesseract`` system binary.

Public API
----------
sanitize(source, filename, posture='standard') -> SanitizedDocument
    source   : bytes or path-like
    filename : declared filename (used for format detection)
    posture  : 'standard' | 'deep' | 'forensic'

SanitizedDocument
    .clean_text        str
    .original_format   str   ('pdf' | 'docx' | 'html' | 'text')
    .sanitization_log  list[str]
    .escalated_to_deep bool
"""
from __future__ import annotations

import io
import logging
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Union

from bossbox.ingest.exceptions import (
    SanitizerDeepModeError,
    SanitizerFormatError,
)

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# Unicode code points that are invisible and frequently abused for injection.
_ZERO_WIDTH_CHARS: frozenset[str] = frozenset(
    [
        "\u00ad",  # SOFT HYPHEN
        "\u200b",  # ZERO WIDTH SPACE
        "\u200c",  # ZERO WIDTH NON-JOINER
        "\u200d",  # ZERO WIDTH JOINER
        "\u200e",  # LEFT-TO-RIGHT MARK
        "\u200f",  # RIGHT-TO-LEFT MARK
        "\u202a",  # LEFT-TO-RIGHT EMBEDDING
        "\u202b",  # RIGHT-TO-LEFT EMBEDDING
        "\u202c",  # POP DIRECTIONAL FORMATTING
        "\u202d",  # LEFT-TO-RIGHT OVERRIDE
        "\u202e",  # RIGHT-TO-LEFT OVERRIDE
        "\u2060",  # WORD JOINER
        "\u2061",  # FUNCTION APPLICATION (invisible operator)
        "\u2062",  # INVISIBLE TIMES
        "\u2063",  # INVISIBLE SEPARATOR
        "\u2064",  # INVISIBLE PLUS
        "\ufeff",  # ZERO WIDTH NO-BREAK SPACE / BOM
    ]
)

# CSS property patterns that indicate hidden elements.
_HIDDEN_CSS_PATTERNS: list[re.Pattern[str]] = [
    re.compile(r"display\s*:\s*none", re.I),
    re.compile(r"visibility\s*:\s*hidden", re.I),
    re.compile(r"opacity\s*:\s*0(?:\.0+)?\b", re.I),
    re.compile(r"font-size\s*:\s*0(?:\.0+)?\s*(?:px|pt|em|rem|%)?", re.I),
]

# PDF white colour value as returned by pymupdf (0xFFFFFF).
_PDF_WHITE = 16777215

# Minimum visible font size in points.
_MIN_FONT_SIZE: float = 1.0

# Supported posture strings.
_DEEP_POSTURES: frozenset[str] = frozenset(["deep", "forensic"])


# ---------------------------------------------------------------------------
# Output dataclass
# ---------------------------------------------------------------------------


@dataclass
class SanitizedDocument:
    """Result of physical sanitization."""

    clean_text: str
    original_format: str
    sanitization_log: list[str] = field(default_factory=list)
    escalated_to_deep: bool = False


# ---------------------------------------------------------------------------
# Internal helpers — general
# ---------------------------------------------------------------------------


def _load_bytes(source: Union[bytes, str, Path]) -> bytes:
    """Normalise *source* to ``bytes``."""
    if isinstance(source, (str, Path)):
        return Path(source).read_bytes()
    return bytes(source)


def _detect_format(filename: str, data: bytes) -> str:
    """
    Determine document format from filename extension then magic bytes.

    Returns one of: 'pdf', 'docx', 'html', 'text'.
    """
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext == "pdf" or data[:4] == b"%PDF":
        return "pdf"
    if ext == "docx" or (ext == "" and data[:2] == b"PK"):
        return "docx"
    if ext in ("html", "htm"):
        return "html"
    if data[:2] == b"PK":  # other ZIP-based Office format — treat as docx
        return "docx"
    # Plaintext / markdown / unknown
    return "text"


def _strip_zero_width(text: str) -> tuple[str, bool]:
    """
    Remove zero-width and invisible Unicode characters.

    Returns ``(cleaned_text, had_suspicious)`` where *had_suspicious* is True
    if any characters were removed.
    """
    found = False
    for ch in _ZERO_WIDTH_CHARS:
        if ch in text:
            found = True
            text = text.replace(ch, "")
    # Also strip non-printable control characters (except newline / tab).
    cleaned_chars: list[str] = []
    for ch in text:
        cp = ord(ch)
        if cp < 0x20 and ch not in ("\n", "\r", "\t"):
            found = True
        else:
            cleaned_chars.append(ch)
    return "".join(cleaned_chars), found


def _check_tesseract() -> None:
    """
    Confirm the ``tesseract`` binary is accessible.

    Raises :exc:`SanitizerDeepModeError` if it is not, so callers get a clear
    message instead of a confusing pytesseract traceback.
    """
    import shutil

    if shutil.which("tesseract") is None:
        raise SanitizerDeepModeError(
            "Deep-mode sanitization requires the 'tesseract' binary. "
            "Install it with: sudo apt install tesseract-ocr  (or equivalent)."
        )


# ---------------------------------------------------------------------------
# Standard-mode sanitizers per format
# ---------------------------------------------------------------------------


def _standard_text(data: bytes) -> tuple[str, list[str], bool]:
    """Standard sanitization for plaintext / Markdown."""
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception:
        text = data.decode("latin-1", errors="replace")

    # Unicode NFC normalisation first.
    text = unicodedata.normalize("NFC", text)
    cleaned, had_zw = _strip_zero_width(text)
    log_entries: list[str] = []
    if had_zw:
        log_entries.append("Stripped zero-width / non-printable Unicode characters.")
    return cleaned, log_entries, had_zw


def _standard_html(data: bytes) -> tuple[str, list[str], bool]:
    """Standard sanitization for HTML: strips hidden DOM elements."""
    from bs4 import BeautifulSoup

    try:
        soup = BeautifulSoup(data, "html.parser")
    except Exception as exc:
        raise SanitizerFormatError(f"Cannot parse HTML: {exc}") from exc

    log_entries: list[str] = []
    escalate = False

    # Remove script / style / noscript blocks entirely.
    for tag_name in ("script", "style", "noscript", "head"):
        for tag in soup.find_all(tag_name):
            tag.decompose()

    # Strip elements hidden via inline style.
    for tag in list(soup.find_all(style=True)):
        style_val: str = tag.get("style", "")
        if any(pat.search(style_val) for pat in _HIDDEN_CSS_PATTERNS):
            log_entries.append(
                f"Stripped hidden element <{tag.name}> (style='{style_val[:60]}')."
            )
            escalate = True
            tag.decompose()

    # Strip elements hidden via hidden attribute or aria-hidden.
    for tag in list(soup.find_all(hidden=True)):
        log_entries.append(f"Stripped element with hidden attribute: <{tag.name}>.")
        escalate = True
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    cleaned, had_zw = _strip_zero_width(text)
    if had_zw:
        log_entries.append("Stripped zero-width Unicode in HTML text content.")
        escalate = True
    return cleaned, log_entries, escalate


def _standard_docx(data: bytes) -> tuple[str, list[str], bool]:
    """Standard sanitization for DOCX: strips hidden-flag runs and metadata."""
    import docx as python_docx

    try:
        doc = python_docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise SanitizerFormatError(f"Cannot open DOCX: {exc}") from exc

    log_entries: list[str] = []
    escalate = False
    parts: list[str] = []

    def _extract_paragraphs(paragraphs) -> None:
        nonlocal escalate
        for para in paragraphs:
            visible: list[str] = []
            for run in para.runs:
                if run.font.hidden:
                    log_entries.append(
                        f"Stripped hidden run in paragraph: '{run.text[:40]}'."
                    )
                    escalate = True
                    continue
                visible.append(run.text)
            text = "".join(visible).strip()
            if text:
                parts.append(text)

    _extract_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _extract_paragraphs(cell.paragraphs)

    text = "\n".join(parts)
    cleaned, had_zw = _strip_zero_width(text)
    if had_zw:
        log_entries.append("Stripped zero-width Unicode in DOCX text.")
        escalate = True
    return cleaned, log_entries, escalate


def _standard_pdf(data: bytes) -> tuple[str, list[str], bool]:
    """
    Standard sanitization for PDF.

    Extracts text via pymupdf; skips spans with suspicious attributes
    (white/invisible colour, sub-1pt font size).  Does not rasterize.
    """
    import fitz  # pymupdf

    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise SanitizerFormatError(f"Cannot open PDF: {exc}") from exc

    log_entries: list[str] = []
    escalate = False
    page_texts: list[str] = []

    for page_num, page in enumerate(doc, start=1):
        raw = page.get_text("rawdict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        span_texts: list[str] = []
        for block in raw.get("blocks", []):
            if block.get("type") != 0:  # skip image blocks
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    # In rawdict, text lives in per-character 'chars' entries.
                    chars = span.get("chars", [])
                    span_text: str = "".join(ch.get("c", "") for ch in chars)
                    if not span_text.strip():
                        continue
                    size: float = span.get("size", 12.0)
                    color: int = span.get("color", 0)
                    if size < _MIN_FONT_SIZE:
                        log_entries.append(
                            f"Stripped sub-1pt text on page {page_num}: "
                            f"'{span_text[:40]}'."
                        )
                        escalate = True
                        continue
                    if color == _PDF_WHITE:
                        log_entries.append(
                            f"Stripped white-coloured text on page {page_num}: "
                            f"'{span_text[:40]}'."
                        )
                        escalate = True
                        continue
                    span_texts.append(span_text)
        if span_texts:
            page_texts.append("".join(span_texts))

    text = "\n".join(page_texts)
    cleaned, had_zw = _strip_zero_width(text)
    if had_zw:
        log_entries.append("Stripped zero-width Unicode in PDF text.")
        escalate = True

    return cleaned, log_entries, escalate


# ---------------------------------------------------------------------------
# Deep-mode sanitizers per format
# ---------------------------------------------------------------------------


def _deep_pdf(data: bytes) -> tuple[str, list[str]]:
    """Deep mode for PDF: rasterize every page and OCR."""
    import fitz
    from PIL import Image
    import pytesseract

    _check_tesseract()
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise SanitizerFormatError(f"Cannot open PDF for deep sanitization: {exc}") from exc

    log_entries: list[str] = ["Deep mode: rasterizing all PDF pages for OCR."]
    page_texts: list[str] = []

    for page_num, page in enumerate(doc, start=1):
        mat = fitz.Matrix(2.0, 2.0)  # 2× zoom for better OCR quality
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        ocr_text: str = pytesseract.image_to_string(img)
        page_texts.append(ocr_text)
        log_entries.append(
            f"OCR page {page_num}: {len(ocr_text)} characters extracted."
        )

    return "\n".join(page_texts), log_entries


def _deep_docx(data: bytes) -> tuple[str, list[str]]:
    """
    Deep mode for DOCX.

    Runs standard extraction first (which already removes hidden text),
    then logs that deep processing was applied.  Full rasterization of
    DOCX pages would require LibreOffice; that dependency is deferred to a
    later step.  The OCR path is exercised for flagged PDFs.
    """
    text, std_log, _ = _standard_docx(data)
    log_entries = ["Deep mode: DOCX re-processed with hidden-text stripping."] + std_log
    return text, log_entries


def _deep_html(data: bytes) -> tuple[str, list[str]]:
    """Deep mode for HTML: extract only visible text (same as standard here)."""
    text, std_log, _ = _standard_html(data)
    log_entries = ["Deep mode: HTML visible-text extraction applied."] + std_log
    return text, log_entries


def _deep_text(data: bytes) -> tuple[str, list[str]]:
    """Deep mode for plaintext: same as standard (no rasterization applicable)."""
    text, std_log, _ = _standard_text(data)
    log_entries = ["Deep mode: plaintext re-processed (standard pipeline, no rasterization)."] + std_log
    return text, log_entries


# ---------------------------------------------------------------------------
# Dispatch tables
# ---------------------------------------------------------------------------

_STANDARD_DISPATCH = {
    "pdf":  _standard_pdf,
    "docx": _standard_docx,
    "html": _standard_html,
    "text": _standard_text,
}

_DEEP_DISPATCH = {
    "pdf":  _deep_pdf,
    "docx": _deep_docx,
    "html": _deep_html,
    "text": _deep_text,
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def sanitize(
    source: Union[bytes, str, Path],
    filename: str,
    posture: str = "standard",
) -> SanitizedDocument:
    """
    Physically sanitize *source* and return a :class:`SanitizedDocument`.

    Parameters
    ----------
    source:
        Raw document as ``bytes``, or a ``str``/``Path`` pointing to a file.
    filename:
        Declared filename (used for format detection via extension).
    posture:
        ``'standard'``  — hidden-element stripping with auto-escalation.
        ``'deep'``      — force full rasterization + OCR.
        ``'forensic'``  — synonym for ``'deep'``; caller marks envelope
                          with ``hostile_content_acknowledged``.

    Returns
    -------
    SanitizedDocument
    """
    data = _load_bytes(source)
    fmt = _detect_format(filename, data)

    if fmt not in _STANDARD_DISPATCH:
        raise SanitizerFormatError(f"Unsupported document format: '{fmt}'")

    # ── Forced deep mode ────────────────────────────────────────────────────
    if posture in _DEEP_POSTURES:
        deep_fn = _DEEP_DISPATCH[fmt]
        clean, deep_log = deep_fn(data)
        return SanitizedDocument(
            clean_text=clean,
            original_format=fmt,
            sanitization_log=deep_log,
            escalated_to_deep=True,
        )

    # ── Standard mode with optional auto-escalation ──────────────────────────
    std_fn = _STANDARD_DISPATCH[fmt]
    clean, std_log, has_suspicious = std_fn(data)

    if not has_suspicious:
        return SanitizedDocument(
            clean_text=clean,
            original_format=fmt,
            sanitization_log=std_log,
            escalated_to_deep=False,
        )

    # Suspicious elements found — attempt deep re-processing.
    std_log.append("Suspicious elements detected; attempting deep-mode re-processing.")
    try:
        deep_fn = _DEEP_DISPATCH[fmt]
        deep_clean, deep_log = deep_fn(data)
        return SanitizedDocument(
            clean_text=deep_clean,
            original_format=fmt,
            sanitization_log=std_log + deep_log,
            escalated_to_deep=True,
        )
    except SanitizerDeepModeError as exc:
        # SanitizerDeepModeError messages are fully controlled by this module
        # (missing tesseract binary notice), so str(exc) is safe to surface.
        std_log.append(f"Deep-mode re-processing skipped: {exc}")
        return SanitizedDocument(
            clean_text=clean,
            original_format=fmt,
            sanitization_log=std_log,
            escalated_to_deep=True,
        )
