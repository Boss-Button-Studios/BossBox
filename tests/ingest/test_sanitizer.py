"""
Step 9 — Physical Document Sanitizer test suite (pytest)
=========================================================

Test classes
------------
TestSanitizedDocument   — dataclass fields, defaults
TestFormatDetection     — _detect_format extension and magic-byte logic
TestZeroWidthStripping  — _strip_zero_width utility
TestTextSanitizer       — plaintext / markdown standard mode
TestHTMLSanitizer       — display:none, visibility:hidden, opacity:0, hidden attr
TestDOCXSanitizer       — hidden runs, tables, zero-width chars
TestPDFSanitizer        — standard extraction, white text, JS detection
TestDeepMode            — deep mode dispatch, tesseract guard, posture='deep'
TestEscalation          — auto-escalation flag and deep fallback
TestPublicAPI           — sanitize() source types, posture variants, format errors
"""
from __future__ import annotations

import io
import shutil
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from bossbox.ingest.sanitizer import (
    SanitizedDocument,
    _detect_format,
    _strip_zero_width,
    sanitize,
)
from bossbox.ingest.exceptions import SanitizerDeepModeError, SanitizerFormatError


# ---------------------------------------------------------------------------
# Fixture helpers
# ---------------------------------------------------------------------------


def make_pdf(text: str = "Hello PDF", white_text: str = "") -> bytes:
    """Build a minimal in-memory PDF using pymupdf."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), text, fontsize=12)
    if white_text:
        page.insert_text((72, 200), white_text, fontsize=12, color=(1, 1, 1))
    return doc.tobytes()


def make_docx(visible: str = "Hello DOCX", hidden: str = "") -> bytes:
    """Build a minimal in-memory DOCX."""
    import docx as python_docx

    doc = python_docx.Document()
    para = doc.add_paragraph()
    para.add_run(visible)
    if hidden:
        run = para.add_run(hidden)
        run.font.hidden = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


SIMPLE_HTML = b"<html><body><p>Visible text</p></body></html>"
HIDDEN_NONE_HTML = (
    b"<html><body>"
    b"<p>Visible</p>"
    b"<p style='display:none'>Injected hidden</p>"
    b"</body></html>"
)
HIDDEN_VIS_HTML = (
    b"<html><body>"
    b"<p>Visible</p>"
    b"<span style='visibility:hidden'>Secret</span>"
    b"</body></html>"
)
HIDDEN_OPACITY_HTML = (
    b"<html><body>"
    b"<p>Visible</p>"
    b"<div style='opacity:0'>Ghost text</div>"
    b"</body></html>"
)
HIDDEN_ATTR_HTML = (
    b"<html><body>"
    b"<p>Visible</p>"
    b"<p hidden>Hidden attr para</p>"
    b"</body></html>"
)
ZERO_WIDTH_TEXT = "Hello\u200bWorld\u200cTest\ufeff".encode()


# ---------------------------------------------------------------------------
# TestSanitizedDocument
# ---------------------------------------------------------------------------


class TestSanitizedDocument:
    def test_fields_present(self):
        doc = SanitizedDocument(
            clean_text="hello",
            original_format="text",
            sanitization_log=["stripped x"],
            escalated_to_deep=False,
        )
        assert doc.clean_text == "hello"
        assert doc.original_format == "text"
        assert doc.sanitization_log == ["stripped x"]
        assert doc.escalated_to_deep is False

    def test_escalated_default_false(self):
        doc = SanitizedDocument(clean_text="x", original_format="text")
        assert doc.escalated_to_deep is False

    def test_log_default_empty_list(self):
        doc = SanitizedDocument(clean_text="x", original_format="text")
        assert doc.sanitization_log == []

    def test_log_is_mutable(self):
        doc = SanitizedDocument(clean_text="x", original_format="text")
        doc.sanitization_log.append("entry")
        assert len(doc.sanitization_log) == 1


# ---------------------------------------------------------------------------
# TestFormatDetection
# ---------------------------------------------------------------------------


class TestFormatDetection:
    def test_pdf_by_extension(self):
        assert _detect_format("file.pdf", b"") == "pdf"

    def test_pdf_by_magic(self):
        assert _detect_format("file.bin", b"%PDF-1.4") == "pdf"

    def test_docx_by_extension(self):
        assert _detect_format("file.docx", b"") == "docx"

    def test_html_by_extension(self):
        assert _detect_format("page.html", b"") == "html"

    def test_htm_by_extension(self):
        assert _detect_format("page.htm", b"") == "html"

    def test_text_by_extension(self):
        assert _detect_format("notes.txt", b"") == "text"

    def test_markdown_by_extension(self):
        assert _detect_format("readme.md", b"") == "text"

    def test_unknown_extension_defaults_to_text(self):
        assert _detect_format("file.xyz", b"plain text") == "text"

    def test_zip_magic_without_extension_returns_docx(self):
        assert _detect_format("file", b"PK\x03\x04") == "docx"


# ---------------------------------------------------------------------------
# TestZeroWidthStripping
# ---------------------------------------------------------------------------


class TestZeroWidthStripping:
    def test_strips_zero_width_space(self):
        cleaned, found = _strip_zero_width("Hel\u200blo")
        assert "\u200b" not in cleaned
        assert found is True

    def test_strips_bom(self):
        cleaned, found = _strip_zero_width("\ufeffHello")
        assert "\ufeff" not in cleaned
        assert found is True

    def test_strips_multiple_types(self):
        text = "\u200b\u200c\u200dHello\u2060\ufeff"
        cleaned, found = _strip_zero_width(text)
        assert cleaned == "Hello"
        assert found is True

    def test_clean_text_unchanged(self):
        cleaned, found = _strip_zero_width("Hello world\nNew line\t tab")
        assert cleaned == "Hello world\nNew line\t tab"
        assert found is False

    def test_strips_non_printable_controls(self):
        cleaned, found = _strip_zero_width("Hello\x01\x07World")
        assert "\x01" not in cleaned
        assert "\x07" not in cleaned
        assert found is True

    def test_preserves_newline_and_tab(self):
        cleaned, _ = _strip_zero_width("Line1\nLine2\tTabbed")
        assert "\n" in cleaned
        assert "\t" in cleaned


# ---------------------------------------------------------------------------
# TestTextSanitizer
# ---------------------------------------------------------------------------


class TestTextSanitizer:
    def test_plain_text_passthrough(self):
        result = sanitize(b"Hello world", "note.txt")
        assert "Hello world" in result.clean_text
        assert result.original_format == "text"
        assert result.escalated_to_deep is False

    def test_strips_zero_width_chars(self):
        result = sanitize(ZERO_WIDTH_TEXT, "note.txt")
        assert "\u200b" not in result.clean_text
        assert "\u200c" not in result.clean_text
        assert "\ufeff" not in result.clean_text

    def test_escalates_when_zero_width_found(self):
        result = sanitize(ZERO_WIDTH_TEXT, "note.txt")
        assert result.escalated_to_deep is True

    def test_log_entry_when_stripped(self):
        result = sanitize(ZERO_WIDTH_TEXT, "note.txt")
        assert any("zero-width" in entry.lower() for entry in result.sanitization_log)

    def test_no_escalation_clean_text(self):
        result = sanitize(b"Clean text with no hidden content.", "readme.md")
        assert result.escalated_to_deep is False
        assert result.sanitization_log == []

    def test_accepts_path_source(self, tmp_path: Path):
        f = tmp_path / "test.txt"
        f.write_bytes(b"File content here")
        result = sanitize(f, "test.txt")
        assert "File content here" in result.clean_text

    def test_accepts_str_path(self, tmp_path: Path):
        f = tmp_path / "test.txt"
        f.write_bytes(b"String path")
        result = sanitize(str(f), "test.txt")
        assert "String path" in result.clean_text


# ---------------------------------------------------------------------------
# TestHTMLSanitizer
# ---------------------------------------------------------------------------


class TestHTMLSanitizer:
    def test_visible_text_preserved(self):
        result = sanitize(SIMPLE_HTML, "page.html")
        assert "Visible text" in result.clean_text
        assert result.escalated_to_deep is False

    def test_strips_display_none(self):
        result = sanitize(HIDDEN_NONE_HTML, "page.html")
        assert "Injected hidden" not in result.clean_text
        assert "Visible" in result.clean_text
        assert result.escalated_to_deep is True

    def test_strips_visibility_hidden(self):
        result = sanitize(HIDDEN_VIS_HTML, "page.html")
        assert "Secret" not in result.clean_text
        assert result.escalated_to_deep is True

    def test_strips_opacity_zero(self):
        result = sanitize(HIDDEN_OPACITY_HTML, "page.html")
        assert "Ghost text" not in result.clean_text
        assert result.escalated_to_deep is True

    def test_strips_hidden_attribute(self):
        result = sanitize(HIDDEN_ATTR_HTML, "page.html")
        assert "Hidden attr para" not in result.clean_text
        assert result.escalated_to_deep is True

    def test_strips_script_tags(self):
        html = b"<html><body><p>Text</p><script>alert('xss')</script></body></html>"
        result = sanitize(html, "page.html")
        assert "alert" not in result.clean_text

    def test_strips_style_tags(self):
        html = b"<html><head><style>body{display:none}</style></head><body><p>Hi</p></body></html>"
        result = sanitize(html, "page.html")
        assert "display" not in result.clean_text

    def test_log_entry_for_hidden_element(self):
        result = sanitize(HIDDEN_NONE_HTML, "page.html")
        assert any("hidden" in e.lower() for e in result.sanitization_log)

    def test_html_format_detected(self):
        result = sanitize(SIMPLE_HTML, "page.html")
        assert result.original_format == "html"

    def test_zero_width_in_html_triggers_escalation(self):
        html = "<html><body><p>Hello\u200bWorld</p></body></html>".encode()
        result = sanitize(html, "page.html")
        assert result.escalated_to_deep is True


# ---------------------------------------------------------------------------
# TestDOCXSanitizer
# ---------------------------------------------------------------------------


class TestDOCXSanitizer:
    def test_visible_text_preserved(self):
        data = make_docx(visible="Visible paragraph")
        result = sanitize(data, "doc.docx")
        assert "Visible paragraph" in result.clean_text

    def test_hidden_text_stripped(self):
        data = make_docx(visible="Normal", hidden="INJECTED")
        result = sanitize(data, "doc.docx")
        assert "INJECTED" not in result.clean_text
        assert "Normal" in result.clean_text

    def test_escalated_when_hidden_text_found(self):
        data = make_docx(visible="OK", hidden="BAD")
        result = sanitize(data, "doc.docx")
        assert result.escalated_to_deep is True

    def test_no_escalation_clean_docx(self):
        data = make_docx(visible="Clean document")
        result = sanitize(data, "doc.docx")
        assert result.escalated_to_deep is False

    def test_log_entry_for_hidden_run(self):
        data = make_docx(visible="OK", hidden="SECRET")
        result = sanitize(data, "doc.docx")
        assert any("hidden" in e.lower() for e in result.sanitization_log)

    def test_docx_format_detected(self):
        data = make_docx()
        result = sanitize(data, "doc.docx")
        assert result.original_format == "docx"

    def test_zero_width_in_docx_triggers_escalation(self):
        data = make_docx(visible="Hello\u200bWorld")
        result = sanitize(data, "doc.docx")
        assert result.escalated_to_deep is True
        assert "\u200b" not in result.clean_text


# ---------------------------------------------------------------------------
# TestPDFSanitizer
# ---------------------------------------------------------------------------


class TestPDFSanitizer:
    def test_visible_text_extracted(self):
        data = make_pdf(text="Hello PDF world")
        result = sanitize(data, "doc.pdf")
        assert "Hello PDF world" in result.clean_text

    def test_white_text_stripped(self):
        data = make_pdf(text="Visible", white_text="HIDDEN WHITE")
        result = sanitize(data, "doc.pdf")
        assert "HIDDEN WHITE" not in result.clean_text
        assert result.escalated_to_deep is True

    def test_white_text_logs_entry(self):
        data = make_pdf(text="OK", white_text="SECRET")
        result = sanitize(data, "doc.pdf")
        assert any("white" in e.lower() for e in result.sanitization_log)

    def test_pdf_format_detected(self):
        data = make_pdf()
        result = sanitize(data, "doc.pdf")
        assert result.original_format == "pdf"

    def test_clean_pdf_no_escalation(self):
        data = make_pdf(text="Normal visible text only")
        result = sanitize(data, "doc.pdf")
        assert result.escalated_to_deep is False

    def test_invalid_pdf_raises_format_error(self):
        with pytest.raises(SanitizerFormatError):
            sanitize(b"not a pdf at all", "doc.pdf")

    def test_pdf_magic_bytes_detected(self):
        data = make_pdf(text="magic test")
        # filename has no .pdf ext — magic bytes should identify it
        result = sanitize(data, "upload.bin")
        assert result.original_format == "pdf"
        assert "magic test" in result.clean_text


# ---------------------------------------------------------------------------
# TestDeepMode
# ---------------------------------------------------------------------------


class TestDeepMode:
    def test_deep_text_same_as_standard(self):
        result = sanitize(b"Deep text content", "note.txt", posture="deep")
        assert "Deep text content" in result.clean_text
        assert result.escalated_to_deep is True

    def test_forensic_posture_sets_flag(self):
        result = sanitize(b"Forensic content", "note.txt", posture="forensic")
        assert result.escalated_to_deep is True

    def test_deep_html_strips_hidden(self):
        result = sanitize(HIDDEN_NONE_HTML, "page.html", posture="deep")
        assert "Injected hidden" not in result.clean_text
        assert result.escalated_to_deep is True

    def test_deep_docx_strips_hidden(self):
        data = make_docx(visible="OK", hidden="BAD")
        result = sanitize(data, "doc.docx", posture="deep")
        assert "BAD" not in result.clean_text
        assert result.escalated_to_deep is True

    def test_deep_pdf_raises_without_tesseract(self):
        data = make_pdf(text="Deep PDF")
        with patch("shutil.which", return_value=None):
            with pytest.raises(SanitizerDeepModeError):
                sanitize(data, "doc.pdf", posture="deep")

    def test_deep_pdf_with_mocked_tesseract(self):
        data = make_pdf(text="OCR content")
        with patch("shutil.which", return_value="/usr/bin/tesseract"), \
             patch("pytesseract.image_to_string", return_value="OCR extracted text"):
            result = sanitize(data, "doc.pdf", posture="deep")
        assert result.clean_text == "OCR extracted text"
        assert result.escalated_to_deep is True

    def test_deep_log_includes_mode_notice(self):
        result = sanitize(b"text content", "note.txt", posture="deep")
        assert any("deep mode" in e.lower() for e in result.sanitization_log)


# ---------------------------------------------------------------------------
# TestEscalation
# ---------------------------------------------------------------------------


class TestEscalation:
    def test_escalated_flag_true_when_suspicious(self):
        data = make_docx(visible="OK", hidden="BAD")
        result = sanitize(data, "doc.docx")
        assert result.escalated_to_deep is True

    def test_escalated_flag_false_when_clean(self):
        result = sanitize(b"Plain clean text", "notes.txt")
        assert result.escalated_to_deep is False

    def test_pdf_escalation_skips_deep_when_no_tesseract(self):
        # A PDF with white text triggers escalation; without tesseract the
        # escalation is recorded but clean_text comes from standard extraction.
        data = make_pdf(text="Visible", white_text="Hidden")
        with patch("shutil.which", return_value=None):
            result = sanitize(data, "doc.pdf")
        assert result.escalated_to_deep is True
        assert any("skipped" in e.lower() for e in result.sanitization_log)

    def test_pdf_escalation_runs_deep_when_tesseract_present(self):
        data = make_pdf(text="Visible", white_text="Hidden")
        with patch("shutil.which", return_value="/usr/bin/tesseract"), \
             patch("pytesseract.image_to_string", return_value="OCR result"):
            result = sanitize(data, "doc.pdf")
        assert result.escalated_to_deep is True
        assert "OCR result" in result.clean_text

    def test_html_escalation_no_deep_needed(self):
        # HTML deep mode doesn't need tesseract — should not raise.
        result = sanitize(HIDDEN_NONE_HTML, "page.html")
        assert result.escalated_to_deep is True
        assert "Injected hidden" not in result.clean_text

    def test_escalation_log_mentions_suspicious_detection(self):
        data = make_docx(visible="OK", hidden="BAD")
        result = sanitize(data, "doc.docx")
        assert any("suspicious" in e.lower() for e in result.sanitization_log)


# ---------------------------------------------------------------------------
# TestPublicAPI
# ---------------------------------------------------------------------------


class TestPublicAPI:
    def test_returns_sanitized_document_instance(self):
        result = sanitize(b"hello", "note.txt")
        assert isinstance(result, SanitizedDocument)

    def test_accepts_bytes(self):
        result = sanitize(b"bytes input", "note.txt")
        assert "bytes input" in result.clean_text

    def test_accepts_path_object(self, tmp_path: Path):
        f = tmp_path / "doc.txt"
        f.write_bytes(b"path input")
        result = sanitize(f, "doc.txt")
        assert "path input" in result.clean_text

    def test_accepts_string_path(self, tmp_path: Path):
        f = tmp_path / "doc.txt"
        f.write_bytes(b"str path input")
        result = sanitize(str(f), "doc.txt")
        assert "str path input" in result.clean_text

    def test_default_posture_is_standard(self):
        result = sanitize(b"test", "note.txt")
        # Standard mode on clean text should NOT force escalated_to_deep
        assert result.escalated_to_deep is False

    def test_sanitization_log_is_list(self):
        result = sanitize(b"test", "note.txt")
        assert isinstance(result.sanitization_log, list)

    def test_clean_text_is_str(self):
        result = sanitize(b"test", "note.txt")
        assert isinstance(result.clean_text, str)

    def test_original_format_is_str(self):
        result = sanitize(b"test", "note.txt")
        assert isinstance(result.original_format, str)
