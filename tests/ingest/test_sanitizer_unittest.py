"""
Step 9 — Physical Document Sanitizer test suite (stdlib unittest)
==================================================================
Identical coverage to test_sanitizer.py; runnable without pytest.

    python -m unittest tests.ingest.test_sanitizer_unittest -v
"""
from __future__ import annotations

import io
import unittest
from pathlib import Path
from unittest.mock import patch


from bossbox.ingest.sanitizer import (
    SanitizedDocument,
    _detect_format,
    _strip_zero_width,
    sanitize,
)
from bossbox.ingest.exceptions import SanitizerDeepModeError, SanitizerFormatError


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------


def _make_pdf(text: str = "Hello PDF", white_text: str = "") -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), text, fontsize=12)
    if white_text:
        page.insert_text((72, 200), white_text, fontsize=12, color=(1, 1, 1))
    return doc.tobytes()


def _make_docx(visible: str = "Hello DOCX", hidden: str = "") -> bytes:
    import docx as python_docx

    doc = python_docx.Document()
    para = doc.add_paragraph()
    para.add_run(visible)
    if hidden:
        run = para.add_run(hidden)
        run.font.hidden = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_SIMPLE_HTML = b"<html><body><p>Visible text</p></body></html>"
_HIDDEN_NONE_HTML = (
    b"<html><body>"
    b"<p>Visible</p>"
    b"<p style='display:none'>Injected hidden</p>"
    b"</body></html>"
)
_HIDDEN_VIS_HTML = (
    b"<html><body>"
    b"<p>Visible</p>"
    b"<span style='visibility:hidden'>Secret</span>"
    b"</body></html>"
)
_ZERO_WIDTH_TEXT = "Hello\u200bWorld\u200cTest\ufeff".encode()


# ---------------------------------------------------------------------------
# SanitizedDocument
# ---------------------------------------------------------------------------


class TestSanitizedDocument(unittest.TestCase):
    def test_fields(self):
        doc = SanitizedDocument(
            clean_text="hi", original_format="text",
            sanitization_log=["x"], escalated_to_deep=True,
        )
        self.assertEqual(doc.clean_text, "hi")
        self.assertEqual(doc.original_format, "text")
        self.assertEqual(doc.sanitization_log, ["x"])
        self.assertTrue(doc.escalated_to_deep)

    def test_defaults(self):
        doc = SanitizedDocument(clean_text="x", original_format="text")
        self.assertFalse(doc.escalated_to_deep)
        self.assertEqual(doc.sanitization_log, [])


# ---------------------------------------------------------------------------
# Format detection
# ---------------------------------------------------------------------------


class TestFormatDetection(unittest.TestCase):
    def test_pdf_ext(self):
        self.assertEqual(_detect_format("file.pdf", b""), "pdf")

    def test_pdf_magic(self):
        self.assertEqual(_detect_format("file.bin", b"%PDF-1.4"), "pdf")

    def test_docx_ext(self):
        self.assertEqual(_detect_format("file.docx", b""), "docx")

    def test_html_ext(self):
        self.assertEqual(_detect_format("page.html", b""), "html")

    def test_htm_ext(self):
        self.assertEqual(_detect_format("page.htm", b""), "html")

    def test_text_ext(self):
        self.assertEqual(_detect_format("notes.txt", b""), "text")

    def test_md_ext(self):
        self.assertEqual(_detect_format("readme.md", b""), "text")

    def test_unknown_defaults_text(self):
        self.assertEqual(_detect_format("file.xyz", b"plain"), "text")

    def test_zip_magic_docx(self):
        self.assertEqual(_detect_format("file", b"PK\x03\x04"), "docx")


# ---------------------------------------------------------------------------
# Zero-width stripping
# ---------------------------------------------------------------------------


class TestZeroWidthStripping(unittest.TestCase):
    def test_strips_zwsp(self):
        cleaned, found = _strip_zero_width("Hel\u200blo")
        self.assertNotIn("\u200b", cleaned)
        self.assertTrue(found)

    def test_strips_bom(self):
        cleaned, found = _strip_zero_width("\ufeffHello")
        self.assertNotIn("\ufeff", cleaned)
        self.assertTrue(found)

    def test_clean_unchanged(self):
        cleaned, found = _strip_zero_width("Hello\nWorld\t!")
        self.assertEqual(cleaned, "Hello\nWorld\t!")
        self.assertFalse(found)

    def test_preserves_newline_tab(self):
        cleaned, _ = _strip_zero_width("A\nB\tC")
        self.assertIn("\n", cleaned)
        self.assertIn("\t", cleaned)

    def test_strips_control_chars(self):
        cleaned, found = _strip_zero_width("H\x01ello\x07")
        self.assertNotIn("\x01", cleaned)
        self.assertTrue(found)


# ---------------------------------------------------------------------------
# Text sanitizer
# ---------------------------------------------------------------------------


class TestTextSanitizer(unittest.TestCase):
    def test_passthrough(self):
        r = sanitize(b"Hello world", "note.txt")
        self.assertIn("Hello world", r.clean_text)
        self.assertFalse(r.escalated_to_deep)

    def test_strips_zero_width(self):
        r = sanitize(_ZERO_WIDTH_TEXT, "note.txt")
        self.assertNotIn("\u200b", r.clean_text)
        self.assertTrue(r.escalated_to_deep)

    def test_log_entry_on_strip(self):
        r = sanitize(_ZERO_WIDTH_TEXT, "note.txt")
        self.assertTrue(any("zero-width" in e.lower() for e in r.sanitization_log))

    def test_no_escalation_clean(self):
        r = sanitize(b"Clean text", "readme.md")
        self.assertFalse(r.escalated_to_deep)

    def test_path_source(self):
        import tempfile, os
        with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as f:
            f.write(b"From file")
            name = f.name
        try:
            r = sanitize(Path(name), "test.txt")
            self.assertIn("From file", r.clean_text)
        finally:
            os.unlink(name)


# ---------------------------------------------------------------------------
# HTML sanitizer
# ---------------------------------------------------------------------------


class TestHTMLSanitizer(unittest.TestCase):
    def test_visible_preserved(self):
        r = sanitize(_SIMPLE_HTML, "page.html")
        self.assertIn("Visible text", r.clean_text)
        self.assertFalse(r.escalated_to_deep)

    def test_display_none_stripped(self):
        r = sanitize(_HIDDEN_NONE_HTML, "page.html")
        self.assertNotIn("Injected hidden", r.clean_text)
        self.assertTrue(r.escalated_to_deep)

    def test_visibility_hidden_stripped(self):
        r = sanitize(_HIDDEN_VIS_HTML, "page.html")
        self.assertNotIn("Secret", r.clean_text)
        self.assertTrue(r.escalated_to_deep)

    def test_strips_scripts(self):
        html = b"<html><body><p>Text</p><script>evil()</script></body></html>"
        r = sanitize(html, "page.html")
        self.assertNotIn("evil", r.clean_text)

    def test_log_entry_for_hidden(self):
        r = sanitize(_HIDDEN_NONE_HTML, "page.html")
        self.assertTrue(any("hidden" in e.lower() for e in r.sanitization_log))

    def test_format_is_html(self):
        r = sanitize(_SIMPLE_HTML, "page.html")
        self.assertEqual(r.original_format, "html")


# ---------------------------------------------------------------------------
# DOCX sanitizer
# ---------------------------------------------------------------------------


class TestDOCXSanitizer(unittest.TestCase):
    def test_visible_preserved(self):
        data = _make_docx(visible="Visible para")
        r = sanitize(data, "doc.docx")
        self.assertIn("Visible para", r.clean_text)

    def test_hidden_stripped(self):
        data = _make_docx(visible="Normal", hidden="INJECTED")
        r = sanitize(data, "doc.docx")
        self.assertNotIn("INJECTED", r.clean_text)
        self.assertIn("Normal", r.clean_text)

    def test_escalated_when_hidden(self):
        data = _make_docx(visible="OK", hidden="BAD")
        r = sanitize(data, "doc.docx")
        self.assertTrue(r.escalated_to_deep)

    def test_no_escalation_clean_docx(self):
        data = _make_docx(visible="Clean")
        r = sanitize(data, "doc.docx")
        self.assertFalse(r.escalated_to_deep)

    def test_log_entry_for_hidden_run(self):
        data = _make_docx(visible="OK", hidden="SECRET")
        r = sanitize(data, "doc.docx")
        self.assertTrue(any("hidden" in e.lower() for e in r.sanitization_log))

    def test_format_is_docx(self):
        r = sanitize(_make_docx(), "doc.docx")
        self.assertEqual(r.original_format, "docx")


# ---------------------------------------------------------------------------
# PDF sanitizer
# ---------------------------------------------------------------------------


class TestPDFSanitizer(unittest.TestCase):
    def test_text_extracted(self):
        data = _make_pdf(text="Hello PDF world")
        r = sanitize(data, "doc.pdf")
        self.assertIn("Hello PDF world", r.clean_text)

    def test_white_text_stripped(self):
        data = _make_pdf(text="Visible", white_text="HIDDEN WHITE")
        r = sanitize(data, "doc.pdf")
        self.assertNotIn("HIDDEN WHITE", r.clean_text)
        self.assertTrue(r.escalated_to_deep)

    def test_clean_pdf_no_escalation(self):
        data = _make_pdf(text="Normal text only")
        r = sanitize(data, "doc.pdf")
        self.assertFalse(r.escalated_to_deep)

    def test_invalid_pdf_raises(self):
        with self.assertRaises(SanitizerFormatError):
            sanitize(b"not a pdf", "doc.pdf")

    def test_format_is_pdf(self):
        data = _make_pdf()
        r = sanitize(data, "doc.pdf")
        self.assertEqual(r.original_format, "pdf")


# ---------------------------------------------------------------------------
# Deep mode
# ---------------------------------------------------------------------------


class TestDeepMode(unittest.TestCase):
    def test_deep_text(self):
        r = sanitize(b"Deep text", "note.txt", posture="deep")
        self.assertIn("Deep text", r.clean_text)
        self.assertTrue(r.escalated_to_deep)

    def test_forensic_posture(self):
        r = sanitize(b"Forensic", "note.txt", posture="forensic")
        self.assertTrue(r.escalated_to_deep)

    def test_deep_html_strips_hidden(self):
        r = sanitize(_HIDDEN_NONE_HTML, "page.html", posture="deep")
        self.assertNotIn("Injected hidden", r.clean_text)
        self.assertTrue(r.escalated_to_deep)

    def test_deep_pdf_raises_without_tesseract(self):
        data = _make_pdf(text="Deep PDF")
        with patch("shutil.which", return_value=None):
            with self.assertRaises(SanitizerDeepModeError):
                sanitize(data, "doc.pdf", posture="deep")

    def test_deep_pdf_with_mocked_tesseract(self):
        data = _make_pdf(text="OCR input")
        with patch("shutil.which", return_value="/usr/bin/tesseract"), \
             patch("pytesseract.image_to_string", return_value="OCR output"):
            r = sanitize(data, "doc.pdf", posture="deep")
        self.assertEqual(r.clean_text, "OCR output")
        self.assertTrue(r.escalated_to_deep)

    def test_deep_log_has_mode_notice(self):
        r = sanitize(b"text", "note.txt", posture="deep")
        self.assertTrue(any("deep mode" in e.lower() for e in r.sanitization_log))


# ---------------------------------------------------------------------------
# Escalation
# ---------------------------------------------------------------------------


class TestEscalation(unittest.TestCase):
    def test_flag_true_on_suspicious(self):
        data = _make_docx(visible="OK", hidden="BAD")
        r = sanitize(data, "doc.docx")
        self.assertTrue(r.escalated_to_deep)

    def test_flag_false_on_clean(self):
        r = sanitize(b"Clean", "note.txt")
        self.assertFalse(r.escalated_to_deep)

    def test_pdf_escalation_skips_deep_no_tesseract(self):
        data = _make_pdf(text="OK", white_text="HIDDEN")
        with patch("shutil.which", return_value=None):
            r = sanitize(data, "doc.pdf")
        self.assertTrue(r.escalated_to_deep)
        self.assertTrue(any("skipped" in e.lower() for e in r.sanitization_log))

    def test_pdf_escalation_runs_deep_with_tesseract(self):
        data = _make_pdf(text="OK", white_text="HIDDEN")
        with patch("shutil.which", return_value="/usr/bin/tesseract"), \
             patch("pytesseract.image_to_string", return_value="OCR result"):
            r = sanitize(data, "doc.pdf")
        self.assertTrue(r.escalated_to_deep)
        self.assertIn("OCR result", r.clean_text)

    def test_escalation_log_mentions_suspicious(self):
        data = _make_docx(visible="OK", hidden="BAD")
        r = sanitize(data, "doc.docx")
        self.assertTrue(any("suspicious" in e.lower() for e in r.sanitization_log))


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


class TestPublicAPI(unittest.TestCase):
    def test_returns_sanitized_document(self):
        self.assertIsInstance(sanitize(b"x", "note.txt"), SanitizedDocument)

    def test_clean_text_is_str(self):
        self.assertIsInstance(sanitize(b"x", "note.txt").clean_text, str)

    def test_original_format_is_str(self):
        self.assertIsInstance(sanitize(b"x", "note.txt").original_format, str)

    def test_log_is_list(self):
        self.assertIsInstance(sanitize(b"x", "note.txt").sanitization_log, list)

    def test_default_posture_standard(self):
        r = sanitize(b"clean", "note.txt")
        self.assertFalse(r.escalated_to_deep)


if __name__ == "__main__":
    unittest.main()
